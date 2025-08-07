"""
Document processing service for handling PDFs, DOCX, and email documents
"""

import io
import logging
import re
from typing import List
import aiohttp
import PyPDF2
from docx import Document
from email_reply_parser import EmailReplyParser

from src.core.config import get_settings
from src.core.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Service for processing various document types"""
    
    def __init__(self):
        self.max_size_bytes = settings.max_doc_size_mb * 1024 * 1024
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
    
    async def process_document_url(self, url: str) -> str:
        """
        Download and process a document from URL
        
        Args:
            url: URL to the document
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            logger.info(f"Downloading document from: {url}")
            
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status != 200:
                        raise DocumentProcessingError(f"Failed to download document: HTTP {response.status}")
                    
                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > self.max_size_bytes:
                        raise DocumentProcessingError(f"Document too large: {content_length} bytes")
                    
                    content = await response.read()
                    
            # Determine file type from URL or content
            if url.lower().endswith('.pdf'):
                return self._process_pdf(content)
            elif url.lower().endswith(('.docx', '.doc')):
                return self._process_docx(content)
            elif 'email' in url.lower() or self._is_email_content(content):
                return self._process_email(content.decode('utf-8', errors='ignore'))
            else:
                # Try to process as text
                return self._process_text(content)
                
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")
    
    def _process_pdf(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise DocumentProcessingError("No text could be extracted from PDF")
            
            return self._clean_text(text)
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process PDF: {str(e)}")
    
    def _process_docx(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise DocumentProcessingError("No text could be extracted from DOCX")
            
            return self._clean_text(text)
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process DOCX: {str(e)}")
    
    def _process_email(self, content: str) -> str:
        """Extract text from email content"""
        try:
            # Use email reply parser to get the main content
            parsed = EmailReplyParser.parse_reply(content)
            
            if not parsed.strip():
                # If parsing fails, return the original content
                parsed = content
            
            return self._clean_text(parsed)
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process email: {str(e)}")
    
    def _process_text(self, content: bytes) -> str:
        """Process plain text content"""
        try:
            text = content.decode('utf-8', errors='ignore')
            return self._clean_text(text)
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process text: {str(e)}")
    
    def _is_email_content(self, content: bytes) -> bool:
        """Check if content appears to be an email"""
        try:
            text = content.decode('utf-8', errors='ignore').lower()
            email_indicators = ['from:', 'to:', 'subject:', 'reply-to:', '@']
            return any(indicator in text for indicator in email_indicators)
        except:
            return False
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\[\]\"\']+', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk = ' '.join(chunk_words)
            
            if chunk.strip():
                chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
